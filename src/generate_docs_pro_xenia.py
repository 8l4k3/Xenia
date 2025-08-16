import markdown
from weasyprint import HTML, CSS
from docx import Document
from pathlib import Path
import qrcode

# === CONFIGURATION ===
fiches_folder = Path("fiches")
qr_folder = Path("qrcode")
qr_folder.mkdir(exist_ok=True)

# Icônes Unicode simples pour sections
ICONS = {
    "allergènes": "🥗",
    "médicaments": "💊",
    "cosmétiques": "💄",
    "scientifique": "🔬",
    "conseil": "💡"
}

# Parcours de toutes les fiches Markdown
for md_file in fiches_folder.glob("*.md"):
    output_pdf = fiches_folder / f"{md_file.stem}.pdf"
    output_docx = fiches_folder / f"{md_file.stem}.docx"
    qr_file = qr_folder / f"{md_file.stem}_QR.png"

    # === 1️⃣ Génération QR code ===
    data_link = f"https://xenia.local/{md_file.name}"  # ou lien relatif
    img = qrcode.make(data_link)
    img.save(qr_file)
    print(f"[OK] QR code généré : {qr_file}")

    # === 2️⃣ Lecture Markdown ===
    with open(md_file, "r", encoding="utf-8") as f:
        md_text = f.read()

    # === 3️⃣ Génération PDF avec mise en page améliorée ===
    html_text = markdown.markdown(md_text)

    html_template = f"""
    <html>
    <head>
    <style>
    body {{
        font-family: Arial, sans-serif;
        margin: 2cm;
        line-height: 1.5;
        color: #333;
    }}
    h1 {{
        color: #0055A4;
        border-bottom: 2px solid #0055A4;
        padding-bottom: 5px;
    }}
    h2 {{
        color: #0077C2;
        margin-top: 1cm;
    }}
    h3 {{
        color: #0099E6;
    }}
    p {{
        margin: 0.5em 0;
    }}
    .section {{
        padding: 0.3cm;
        margin: 0.5cm 0;
        border-left: 4px solid #0055A4;
        background-color: #f0f8ff;
    }}
    .qr {{
        text-align: center;
        margin-bottom: 1cm;
    }}
    </style>
    </head>
    <body>
    <div class="qr">
        <img src="{qr_file}" width="120"/>
        <p>Scanner pour accès en ligne 🔗</p>
    </div>
    {html_text}
    </body>
    </html>
    """
    HTML(string=html_template).write_pdf(output_pdf)
    print(f"[OK] PDF généré : {output_pdf}")

    # === 4️⃣ Génération DOCX ===
    doc = Document()
    doc.add_paragraph(f"QR Code pour accès en ligne 🔗 : {qr_file}")

    for line in md_text.splitlines():
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            # Détection simple de mots clés pour icônes
            paragraph_text = line
            for key, icon in ICONS.items():
                if key.lower() in line.lower():
                    paragraph_text = f"{icon} {line}"
                    break
            doc.add_paragraph(paragraph_text)

    doc.save(output_docx)
    print(f"[OK] DOCX généré : {output_docx}")

print("[FIN] Tous les documents générés avec mise en page professionnelle ✅")
